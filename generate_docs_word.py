import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    # Configuración de márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Paleta de colores corporativos
    COLOR_PRIMARY = RGBColor(0, 58, 108)     # #003A6C Azul Corporativo
    COLOR_SECONDARY = RGBColor(13, 148, 136) # #0D9488 Teal
    COLOR_DARK = RGBColor(30, 41, 59)        # #1E293B Slate Dark
    COLOR_MUTED = RGBColor(100, 116, 139)    # #64748B Slate Muted
    COLOR_CODE_BG = "F1F5F9"                 # Fondo gris claro código
    COLOR_BOX_BG = "EEF5FA"                  # Fondo azul muy claro
    COLOR_PLACEHOLDER_BG = "F8FAFC"          # Fondo placeholder capturas

    # Helper para sombrear celdas de tabla
    def set_cell_background(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Helper para márgenes internos de celdas
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Helper para agregar títulos estilizados
    def add_custom_heading(text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run.font.size = Pt(16)
            run.font.color.rgb = COLOR_PRIMARY
            run.font.name = 'Arial'
        elif level == 2:
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run.font.size = Pt(13)
            run.font.color.rgb = COLOR_SECONDARY
            run.font.name = 'Arial'
        elif level == 3:
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_DARK
            run.font.name = 'Arial'
        return p

    # Helper para párrafos estándar
    def add_p(text, bold=False, italic=False, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_DARK
        run.font.name = 'Arial'
        return p

    # Helper para bullet points
    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(9.5)
        run_b.font.color.rgb = COLOR_DARK
        run_b.font.name = 'Arial'
        
        run_t = p.add_run(" " + text)
        run_t.font.size = Pt(9.5)
        run_t.font.color.rgb = COLOR_DARK
        run_t.font.name = 'Arial'
        return p

    # Helper para bloques de código
    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, COLOR_CODE_BG)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(15, 23, 42)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Helper para Callout Box (Nota / Importante)
    def add_callout_box(title, text, is_warning=False):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        bg = "FEF2F2" if is_warning else COLOR_BOX_BG
        set_cell_background(cell, bg)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r_title = p.add_run(title + "\n")
        r_title.bold = True
        r_title.font.name = 'Arial'
        r_title.font.size = Pt(9.5)
        r_title.font.color.rgb = RGBColor(185, 28, 28) if is_warning else COLOR_PRIMARY
        
        r_text = p.add_run(text)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(9)
        r_text.font.color.rgb = COLOR_DARK
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Helper para Cajas de Capturas de Pantalla (Placeholder con marco)
    def add_screenshot_placeholder(caption, description):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, COLOR_PLACEHOLDER_BG)
        set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        
        run_icon = p.add_run("📸 [PEGAR AQUÍ CAPTURA DE PANTALLA]\n")
        run_icon.bold = True
        run_icon.font.name = 'Arial'
        run_icon.font.size = Pt(11)
        run_icon.font.color.rgb = COLOR_SECONDARY
        
        run_desc = p.add_run(f"({description})")
        run_desc.italic = True
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(9)
        run_desc.font.color.rgb = COLOR_MUTED
        
        # Leyenda de la figura
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run(f"Figura: {caption}")
        r_cap.bold = True
        r_cap.font.name = 'Arial'
        r_cap.font.size = Pt(9)
        r_cap.font.color.rgb = COLOR_MUTED

    # ==========================================
    # PORTADA
    # ==========================================
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Pt(40)
    p_top.paragraph_format.space_after = Pt(10)
    
    r_empresa = p_top.add_run("TECNOFILM S.A. | SIGPLAST ERP")
    r_empresa.bold = True
    r_empresa.font.name = 'Arial'
    r_empresa.font.size = Pt(12)
    r_empresa.font.color.rgb = COLOR_SECONDARY

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(12)
    r_title = p_title.add_run("SISTEMA WEB DE TRANSFERENCIA DE ALMACENES")
    r_title.bold = True
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(24)
    r_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(30)
    r_sub = p_sub.add_run("Documentación Técnica Integral, Manual de Usuario y Guía de Despliegue Local")
    r_sub.italic = True
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_MUTED

    # Tabla resumen de metadatos de portada
    tbl_meta = doc.add_table(rows=5, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Proyecto:", "Sistema Web de Transferencia de Almacenes (SIGPLAST)"),
        ("Cliente / Empresa:", "Tecnofilm"),
        ("Stack Tecnológico:", "Angular 22, Tailwind CSS 4.3, PHP 8.2 Native, SQL Server 2022"),
        ("Versión del Sistema:", "1.0.0 (Producción / Local)"),
        ("Fecha de Emisión:", "Febrero 2026")
    ]
    for idx, (label, val) in enumerate(meta_data):
        c0 = tbl_meta.cell(idx, 0)
        c1 = tbl_meta.cell(idx, 1)
        set_cell_background(c0, "F8FAFC")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, 80, 80, 100, 100)
        set_cell_margins(c1, 80, 80, 100, 100)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        r0.font.name = 'Arial'
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = COLOR_PRIMARY
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.name = 'Arial'
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_DARK

    doc.add_page_break()

    # ==========================================
    # SECCIÓN 1: INTRODUCCIÓN Y RESUMEN EJECUTIVO
    # ==========================================
    add_custom_heading("1. Introducción y Resumen Ejecutivo", level=1)
    add_p("El presente documento detalla la arquitectura técnica, modelo de datos, ingeniería inversa criptográfica, procedimiento almacenado optimizado, manual de usuario y guía de puesta en marcha del Sistema Web Independiente para la Transferencia de Almacenes desarrollado para Tecnofilm.")
    
    add_custom_heading("1.1. Diagnóstico del Problema", level=2)
    add_p("El ERP central de la empresa Tecnofilm opera sobre una infraestructura cliente-servidor basada en una aplicación de escritorio Windows. Dicho esquema presentaba las siguientes restricciones operativas:")
    add_bullet("Dependencia de Terminales Fijos:", "Los operarios de almacén y planta debían trasladarse físicamente a terminales de escritorio para consultar existencias, emitir vales o validar movimientos.")
    add_bullet("Retrasos en Captura de Datos:", "La necesidad de apuntar transacciones en papel para luego digitarlas generaba discrepancias temporales en el stock real.")
    add_bullet("Falta de Movilidad:", "No existía una interfaz adaptada a dispositivos móviles, tablets o pistolas colectoras de datos (Handhelds / PDA) para operar a pie de planta.")

    add_custom_heading("1.2. Solución Implementada", level=2)
    add_p("Se diseñó y construyó un sistema web progresivo (SPA) desacoplado, ligero y responsivo que replica con total fidelidad la lógica de transferencias y notas de salida del ERP de escritorio. La solución se comunica directamente con la base de datos central de SQL Server 2022 sin alterar las tablas ni procedimientos de la empresa.")

    # ==========================================
    # SECCIÓN 2: STACK TECNOLÓGICO Y ARQUITECTURA
    # ==========================================
    add_custom_heading("2. Stack Tecnológico y Arquitectura del Sistema", level=1)
    add_p("El sistema está construido bajo una arquitectura cliente-servidor desacoplada mediante APIs REST en formato JSON:")

    # Tabla Stack
    tbl_tech = doc.add_table(rows=6, cols=3)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    tech_headers = ["Capa", "Tecnología / Versión", "Rol y Justificación Técnica"]
    for i, h in enumerate(tech_headers):
        c = tbl_tech.cell(0, i)
        set_cell_background(c, "003A6C")
        set_cell_margins(c, 100, 100, 120, 120)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Arial'
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    tech_rows = [
        ("Frontend", "Angular 22.0.0", "Standalone Components, Signals reactivas, Router y ChangeDetectionStrategy OnPush."),
        ("Estilos UI", "Tailwind CSS 4.3.3", "Diseño utilitario moderno, dark mode, personalización de scrollbars y mobile-first."),
        ("Notificaciones", "SweetAlert2 11.26.25", "Modales y confirmaciones asíncronas visuales adaptadas a temas claro y oscuro."),
        ("Backend API", "PHP 8.2 Nativo + XAMPP", "Micro-endpoints REST ligeros con procesamiento nativo de JSON y driver sqlsrv."),
        ("Base de Datos", "Microsoft SQL Server 2022", "Base de datos corporativa TECNOTEST, tablas maestras y procedimientos almacenados ACID.")
    ]
    for idx, (c1, c2, c3) in enumerate(tech_rows):
        row_cells = tbl_tech.rows[idx + 1].cells
        bg_c = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate([c1, c2, c3]):
            cell = row_cells[i]
            set_cell_background(cell, bg_c)
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_DARK
            if i == 0:
                r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ==========================================
    # SECCIÓN 3: BASE DE DATOS HEREDADA (SQL SERVER)
    # ==========================================
    add_custom_heading("3. Modelo de Datos Heredado (SQL Server 2022)", level=1)
    add_p("El sistema opera sobre la base de datos corporativa TECNOTEST, interactuando con las siguientes tablas clave:")
    add_bullet("adm_usuario:", "Almacena credenciales de usuarios (usr_codigo, usr_clave, usr_status, pus_codigo).")
    add_bullet("mae_almacen:", "Catálogo de almacenes físicos (001: Almacén Verde MP, 002: Almacén MP 02, 016: Almacén 016).")
    add_bullet("mae_articulo:", "Maestro de artículos, códigos EAN/barras, descripciones y unidades de medida (ume_codigo).")
    add_bullet("mae_ccostomae:", "Catálogo de centros de costos asignables a cada detalle de movimiento (cco_status = 1).")
    add_bullet("mae_tpomov:", "Tipos de movimiento de almacén (102: Salida por Transferencia, 023: Ingreso por Transferencia).")
    add_bullet("mae_correlativo:", "Control de numeración autoincremental de vales por almacén, año y tipo 'MOV'.")
    add_bullet("log_cabmov:", "Cabecera transaccional de movimientos (mov_id, mov_codigo, mov_tipo, fechas, glosa).")
    add_bullet("log_detmov:", "Detalle de artículos por movimiento (art_codigo, lot_codigo, mov_ctdmov, cco_codigo, costos).")
    add_bullet("log_stkarticulo:", "Histórico mensual de saldos y valuación de costos unitarios en soles y dólares.")

    # ==========================================
    # SECCIÓN 4: PROCEDIMIENTO ALMACENADO MEJORADO
    # ==========================================
    add_custom_heading("4. Procedimiento Almacenado: sp_registrar_transferencia_almacen (Versión Optimizada)", level=1)
    add_p("Ubicación: backend/api/erp/sp_update_v3.sql")
    add_p("Para garantizar la máxima integridad transaccional (ACID), prevenir problemas de concurrencia y optimizar los tiempos de respuesta en el servidor SQL Server 2022, se implementó una versión refactorizada y robustecida del procedimiento almacenado sp_registrar_transferencia_almacen.")
    
    add_custom_heading("4.1. Cambios y Mejoras Técnicas Aplicadas", level=2)
    add_bullet("1. Bloqueo Exclusivo Controlado (UPDLOCK):", "Se implementó el hint WITH (UPDLOCK) en la instrucción MERGE sobre mae_correlativo. Esto previene interbloqueos (deadlocks) en momentos de alta concurrencia, asegurando que la transacción que lee el correlativo sea la única que lo incremente.")
    add_bullet("2. Optimización en Validación de Stock (CTE PedidoAgrupado):", "Se reemplazaron las subconsultas pesadas anidadas por una Expresión de Tabla Común (CTE). Los artículos pedidos se agrupan una sola vez y se calculan contra el saldo físico real mediante OUTER APPLY, reduciendo el costo de CPU y lecturas lógicas.")
    add_bullet("3. Inserción Masiva Simplificada (Principio DRY con #TempDetalle):", "En lugar de procesar el JSON dos veces idénticas para la salida y el ingreso, los datos parseados y sus costos unitarios (obtenidos de log_stkarticulo) se almacenan en una tabla temporal (#TempDetalle). Esto reduce a la mitad el procesamiento de OPENJSON y lecturas a disco.")
    add_bullet("4. Manejo de Transacciones Seguro (@@TRANCOUNT y Limpieza):", "El bloque CATCH verifica el estado transaccional activo (IF @@TRANCOUNT > 0 ROLLBACK TRANSACTION) y garantiza la eliminación de la tabla temporal si ocurre un error inesperado.")
    add_bullet("5. Formato y Estandarización:", "Estandarización de palabras clave SQL en mayúsculas, asignación directa de variables y manejo claro de mensajes de error.")

    add_custom_heading("4.2. Código SQL Completo del Stored Procedure", level=2)
    add_code_block("""USE [TECNOTEST]
GO
SET ANSI_NULLS ON
GO
SET QUOTED_IDENTIFIER ON
GO

ALTER PROCEDURE [dbo].[sp_registrar_transferencia_almacen]
    @emp_codigo char(3),
    @suc_codigo char(3),
    @alm_origen char(3),
    @alm_destino char(3),
    @fec_emi datetime,
    @usuario varchar(15),
    @detalles_json nvarchar(max),
    @glosa varchar(200)
AS
BEGIN 
    SET NOCOUNT ON;
    SET XACT_ABORT ON;

    BEGIN TRY
        -- =============================================
        -- 1. VALIDACIONES INICIALES BÁSICAS
        -- =============================================
        IF @alm_origen = @alm_destino
        BEGIN
            SELECT 0 AS success, 'El almacen origen y destino no pueden ser el mismo' AS message,
                   NULL AS mov_id_salida, NULL AS vale_salida, NULL AS mov_id_ingreso, NULL AS vale_ingreso;
            RETURN;
        END

        IF ISNULL(@detalles_json, '') = '' OR @detalles_json = '[]'
        BEGIN
            SELECT 0 AS success, 'Debe incluir al menos un articulo para transferir' AS message,
                   NULL AS mov_id_salida, NULL AS vale_salida, NULL AS mov_id_ingreso, NULL AS vale_ingreso;
            RETURN;
        END

        IF EXISTS (
            SELECT 1 
            FROM OPENJSON(@detalles_json) AS item 
            WHERE ISNULL(NULLIF(JSON_VALUE(item.value, '$.lot_codigo'), ''), '') = ''
        )
        BEGIN
            SELECT 0 AS success, 'Todos los artículos deben tener asignado un Lote válido' AS message,
                   NULL AS mov_id_salida, NULL AS vale_salida, NULL AS mov_id_ingreso, NULL AS vale_ingreso;
            RETURN;
        END

        -- =============================================
        -- 2. DECLARACIÓN DE VARIABLES DE ENTORNO
        -- =============================================
        DECLARE @anho_actual varchar(4) = LTRIM(RTRIM(CONVERT(varchar(4), YEAR(@fec_emi))));
        DECLARE @mes_actual varchar(2)  = RIGHT('0' + LTRIM(RTRIM(CONVERT(varchar(2), MONTH(@fec_emi)))), 2);
        DECLARE @anho_2dig varchar(2)   = RIGHT(@anho_actual, 2);
        DECLARE @errores_stock varchar(max);

        -- =============================================
        -- 3. VALIDACIÓN DE STOCK REAL (OPTIMIZADO CON CTE)
        -- =============================================
        -- Sumarizamos lo pedido en el JSON
        WITH PedidoAgrupado AS (
            SELECT 
                JSON_VALUE(item.value, '$.art_codigo') AS art_codigo,
                JSON_VALUE(item.value, '$.art_nombre') AS art_nombre,
                ISNULL(NULLIF(JSON_VALUE(item.value, '$.lot_codigo'), ''), '') AS lot_codigo,
                SUM(CAST(JSON_VALUE(item.value, '$.cantidad') AS DECIMAL(18,3))) AS total_pedida
            FROM OPENJSON(@detalles_json) AS item
            GROUP BY 
                JSON_VALUE(item.value, '$.art_codigo'),
                JSON_VALUE(item.value, '$.art_nombre'),
                ISNULL(NULLIF(JSON_VALUE(item.value, '$.lot_codigo'), ''), '')
        )
        SELECT @errores_stock = STUFF((
            SELECT CHAR(10) + '- ' + ISNULL(p.art_nombre, p.art_codigo) + 
                   ' (Disp: ' + CAST(CAST(ISNULL(stk.stock_real, 0) AS INT) AS VARCHAR) + 
                   ', Sol: ' + CAST(CAST(p.total_pedida AS INT) AS VARCHAR) + ')'
            FROM PedidoAgrupado p
            OUTER APPLY (
                SELECT 
                    SUM(CASE WHEN b.mov_tipo = 'I' THEN a.mov_ctdmov ELSE 0 END) -
                    SUM(CASE WHEN b.mov_tipo = 'S' THEN a.mov_ctdmov ELSE 0 END) AS stock_real
                FROM log_detmov a
                INNER JOIN log_cabmov b ON a.mov_id = b.mov_id
                WHERE b.mov_flag <> 'A' 
                  AND b.alm_codigo = @alm_origen
                  AND b.mov_fecemi <= @fec_emi
                  AND a.art_codigo = p.art_codigo
                  AND ISNULL(a.lot_codigo, '') = p.lot_codigo
            ) stk
            WHERE p.total_pedida > ISNULL(stk.stock_real, 0)
            FOR XML PATH('')
        ), 1, 1, '');

        IF @errores_stock IS NOT NULL AND LEN(@errores_stock) > 0
        BEGIN
            SELECT 0 AS success, 
                   'Stock insuficiente para los siguientes artículos:' + @errores_stock AS message,
                   NULL AS mov_id_salida, NULL AS vale_salida, NULL AS mov_id_ingreso, NULL AS vale_ingreso;
            RETURN;
        END

        -- =============================================
        -- 4. INICIO DE TRANSACCIÓN ATÓMICA
        -- =============================================
        BEGIN TRANSACTION;

        -- 4.1. Generación de Correlativo: SALIDA (Almacén Origen con UPDLOCK)
        DECLARE @num_salida int;
        MERGE mae_correlativo WITH (UPDLOCK) AS target
        USING (SELECT @emp_codigo AS emp, @alm_origen AS alm, 'MOV' AS tipo, @anho_actual AS anho) AS source
        ON target.emp_codigo = source.emp AND target.alm_codigo = source.alm 
           AND target.cor_tipo = source.tipo AND target.cor_anho = source.anho
        WHEN MATCHED THEN
            UPDATE SET cor_numero = RIGHT('0000000' + CAST(CAST(LTRIM(RTRIM(target.cor_numero)) AS INT) + 1 AS VARCHAR(7)), 7)
        WHEN NOT MATCHED THEN
            INSERT (emp_codigo, suc_codigo, alm_codigo, cor_tipo, cor_anho, cor_nmes, cor_serie, cor_numero, cor_desde, cor_hasta)
            VALUES (source.emp, '   ', source.alm, source.tipo, source.anho, '  ', '   ', '0000001', '       ', '       ');

        SELECT @num_salida = CAST(LTRIM(RTRIM(cor_numero)) AS INT)
        FROM mae_correlativo 
        WHERE emp_codigo = @emp_codigo AND alm_codigo = @alm_origen
          AND cor_tipo = 'MOV' AND cor_anho = @anho_actual;

        DECLARE @mov_codigo_salida char(7) = RIGHT('0000000' + CAST(@num_salida AS VARCHAR(7)), 7);
        DECLARE @mov_id_salida char(10) = @anho_2dig + @alm_origen + RIGHT('00000' + CAST(@num_salida AS VARCHAR(5)), 5);

        -- 4.2. Generación de Correlativo: INGRESO (Almacén Destino con UPDLOCK)
        DECLARE @num_ingreso int;
        MERGE mae_correlativo WITH (UPDLOCK) AS target
        USING (SELECT @emp_codigo AS emp, @alm_destino AS alm, 'MOV' AS tipo, @anho_actual AS anho) AS source
        ON target.emp_codigo = source.emp AND target.alm_codigo = source.alm 
           AND target.cor_tipo = source.tipo AND target.cor_anho = source.anho
        WHEN MATCHED THEN
            UPDATE SET cor_numero = RIGHT('0000000' + CAST(CAST(LTRIM(RTRIM(target.cor_numero)) AS INT) + 1 AS VARCHAR(7)), 7)
        WHEN NOT MATCHED THEN
            INSERT (emp_codigo, suc_codigo, alm_codigo, cor_tipo, cor_anho, cor_nmes, cor_serie, cor_numero, cor_desde, cor_hasta)
            VALUES (source.emp, '   ', source.alm, source.tipo, source.anho, '  ', '   ', '0000001', '       ', '       ');

        SELECT @num_ingreso = CAST(LTRIM(RTRIM(cor_numero)) AS INT)
        FROM mae_correlativo 
        WHERE emp_codigo = @emp_codigo AND alm_codigo = @alm_destino
          AND cor_tipo = 'MOV' AND cor_anho = @anho_actual;

        DECLARE @mov_codigo_ingreso char(7) = RIGHT('0000000' + CAST(@num_ingreso AS VARCHAR(7)), 7);
        DECLARE @mov_id_ingreso char(10) = @anho_2dig + @alm_destino + RIGHT('00000' + CAST(@num_ingreso AS VARCHAR(5)), 5);

        -- =============================================
        -- 5. INSERCIÓN DE CABECERAS (log_cabmov)
        -- =============================================
        -- Cabecera Salida (102)
        INSERT INTO log_cabmov (
            emp_codigo, mov_id, suc_codigo, alm_codigo, mov_codigo, cli_codigo, cls_codigo, cca_codigo, mov_tipo, mov_anho,
            mov_nmes, tmo_codigo, mov_fecemi, mov_fectras, mov_fecreg, ord_id, opr_id, mov_tpdoc, mov_serdoc, mov_nrodoc,
            mov_fserdoc, mov_fnrodoc, mov_indgre, codtrans, veh_placa, cho_codigo, mov_lugent, mov_glosa, mov_print, mov_flag,
            id_venta, cde_numdes, mov_inddev, mov_rectra, mov_indtra, mov_indcod, mov_idref, mre_codigo, mov_almdes, dsp_id,
            mov_indrnd, mov_indave, prv_codigo, mov_indsor, mov_indcom, mov_indapc, mov_dspspe, ncompra, user_, date,
            time, mov_indrdl, mov_indavc, edp_codigo, acc_id, numped, cco_codigo, rpr_codigo, req_id, mov_anufac,
            mov_movapr, mov_reqapr, mov_deplen, mov_pvnlen, mov_dislen, pca_numero, mov_proser, mov_prodoc, mov_proimp, prp_codigo,
            etp_codigo, user_imp, date_imp, time_imp, user_anu, date_anu, time_anu, ordcli, pesaje, transp,
            ructrans, licencia, placa, unidad, const_ins, cho_nombre, dni, envio_sunat, transporte, nro_bultos,
            peso_bruto, respuesta_nubefact, cvt_codigo, observa, destino, ruc_destino, factor_imp, mon_codigo
        ) VALUES (
            @emp_codigo, @mov_id_salida, @suc_codigo, @alm_origen, @mov_codigo_salida, '000001', '', '', 'S', @anho_actual,
            @mes_actual, '102', @fec_emi, @fec_emi, '19000101', '', '', '-', '', '',
            '', '', 0, '', '', '', '', @glosa, 0, ' ',
            '', '', 0, 0, 1, 0, @mov_id_ingreso, '', @alm_destino, '',
            0, 0, '', 0, 0, 0, 0, '', @usuario, GETDATE(),
            CONVERT(VARCHAR(8), GETDATE(), 108), 0, 0, '', '', '', '', '', '', 0,
            0, 0, '', '', '', '', '', '', 0, '',
            '', '', '19000101', '', '', '19000101', '', '', '', '',
            '', '', '', '', '', '', '', '', '', 0,
            0.00, '', '', '', '', '', 0.00, ''
        );

        -- Cabecera Ingreso (023)
        INSERT INTO log_cabmov (
            emp_codigo, mov_id, suc_codigo, alm_codigo, mov_codigo, cli_codigo, cls_codigo, cca_codigo, mov_tipo, mov_anho,
            mov_nmes, tmo_codigo, mov_fecemi, mov_fectras, mov_fecreg, ord_id, opr_id, mov_tpdoc, mov_serdoc, mov_nrodoc,
            mov_fserdoc, mov_fnrodoc, mov_indgre, codtrans, veh_placa, cho_codigo, mov_lugent, mov_glosa, mov_print, mov_flag,
            id_venta, cde_numdes, mov_inddev, mov_rectra, mov_indtra, mov_indcod, mov_idref, mre_codigo, mov_almdes, dsp_id,
            mov_indrnd, mov_indave, prv_codigo, mov_indsor, mov_indcom, mov_indapc, mov_dspspe, ncompra, user_, date,
            time, mov_indrdl, mov_indavc, edp_codigo, acc_id, numped, cco_codigo, rpr_codigo, req_id, mov_anufac,
            mov_movapr, mov_reqapr, mov_deplen, mov_pvnlen, mov_dislen, pca_numero, mov_proser, mov_prodoc, mov_proimp, prp_codigo,
            etp_codigo, user_imp, date_imp, time_imp, user_anu, date_anu, time_anu, ordcli, pesaje, transp,
            ructrans, licencia, placa, unidad, const_ins, cho_nombre, dni, envio_sunat, transporte, nro_bultos,
            peso_bruto, respuesta_nubefact, cvt_codigo, observa, destino, ruc_destino, factor_imp, mon_codigo
        ) VALUES (
            @emp_codigo, @mov_id_ingreso, @suc_codigo, @alm_destino, @mov_codigo_ingreso, '000001', '', '', 'I', @anho_actual,
            @mes_actual, '023', @fec_emi, '19000101', '19000101', '', '', '-', '', '',
            '', '', 0, '', '', '', '', @glosa, 0, ' ',
            '', '', 0, 1, 1, 0, @mov_id_salida, '', @alm_origen, '',
            0, 0, '', 0, 0, 0, 0, '', @usuario, GETDATE(),
            CONVERT(VARCHAR(8), GETDATE(), 108), 0, 0, '', '', '', '', '', '', 0,
            0, 0, '', '', '', '', '', '', 0, '',
            '', '', '19000101', '', '', '19000101', '', '', '', '',
            '', '', '', '', '', '', '', '', '', 0,
            0.00, '', '', '', '', '', 0.00, ''
        );

        -- =============================================
        -- 6. PREPARACIÓN E INSERCIÓN DE DETALLES (DRY)
        -- =============================================
        -- 6.1. Parsear JSON y buscar costos una sola vez en tabla temporal
        SELECT 
            JSON_VALUE(item.value, '$.art_codigo') AS art_codigo,
            JSON_VALUE(item.value, '$.art_nombre') AS art_nombre,
            ISNULL(JSON_VALUE(item.value, '$.lot_codigo'), '') AS lot_codigo,
            RIGHT('00' + CAST(CAST(item.[key] AS INT) AS VARCHAR(2)), 2) AS mov_item,
            CAST(JSON_VALUE(item.value, '$.cantidad') AS DECIMAL(18, 3)) AS cantidad,
            ISNULL(JSON_VALUE(item.value, '$.cco_codigo'), '') AS cco_codigo,
            ISNULL(stk.costo_sol, 0.000) AS costo_sol,
            ISNULL(stk.costo_dol, 0.000) AS costo_dol
        INTO #TempDetalle
        FROM OPENJSON(@detalles_json) AS item
        OUTER APPLY (
            SELECT TOP 1 stk_vfin AS costo_sol, stk_vfind AS costo_dol
            FROM log_stkarticulo 
            WHERE emp_codigo = @emp_codigo 
              AND alm_codigo = @alm_origen 
              AND art_codigo = JSON_VALUE(item.value, '$.art_codigo')
              AND stk_anho = @anho_actual 
              AND stk_nmes = @mes_actual
            ORDER BY CASE WHEN lot_codigo = ISNULL(JSON_VALUE(item.value, '$.lot_codigo'), '') THEN 0 ELSE 1 END
        ) stk;

        -- 6.2. Insertar Detalle Salida
        INSERT INTO log_detmov (
            emp_codigo, mov_id, art_codigo, art_nombre, art_proceso, lot_codigo, art_codref,
            mov_item, mov_ctdcom, mov_ctdmov, mov_cuni, mov_ctotal,
            mov_vuni, mov_vpro, mov_vmov, mov_vcom, 
            mov_vunid, mov_vprod, mov_vmovd,
            mov_indcla, mov_indbon, mov_indnc, mov_idnota, mov_vaso, mov_artbon,
            no_ord_tra, mov_sec, kilos, mov_detalle, bobinas, metros, paquetes, millares, merma,
            f_fab, f_ven, cco_codigo
        )
        SELECT 
            @emp_codigo, @mov_id_salida, art_codigo, art_nombre, '', lot_codigo, '',
            mov_item, 0.000, cantidad, 0.000, 0.000,
            costo_sol, costo_sol, costo_sol, 0.000,
            costo_dol, costo_dol, costo_dol,
            0, 0, 0, '', 0, '',
            '', '', 0.000, '', 0.000, 0.00, 0, 0.000, 0.00,
            '19000101', '19000101', cco_codigo
        FROM #TempDetalle;

        -- 6.3. Insertar Detalle Ingreso (usando los mismos datos temporales)
        INSERT INTO log_detmov (
            emp_codigo, mov_id, art_codigo, art_nombre, art_proceso, lot_codigo, art_codref,
            mov_item, mov_ctdcom, mov_ctdmov, mov_cuni, mov_ctotal,
            mov_vuni, mov_vpro, mov_vmov, mov_vcom, 
            mov_vunid, mov_vprod, mov_vmovd,
            mov_indcla, mov_indbon, mov_indnc, mov_idnota, mov_vaso, mov_artbon,
            no_ord_tra, mov_sec, kilos, mov_detalle, bobinas, metros, paquetes, millares, merma,
            f_fab, f_ven, cco_codigo
        )
        SELECT 
            @emp_codigo, @mov_id_ingreso, art_codigo, art_nombre, '', lot_codigo, '',
            mov_item, 0.000, cantidad, 0.000, 0.000,
            costo_sol, costo_sol, costo_sol, 0.000,
            costo_dol, costo_dol, costo_dol,
            0, 0, 0, '', 0, '',
            '', '', 0.000, '', 0.000, 0.00, 0, 0.000, 0.00,
            '19000101', '19000101', cco_codigo
        FROM #TempDetalle;

        -- Limpiar tabla temporal
        DROP TABLE #TempDetalle;

        -- =============================================
        -- 7. COMMIT DE LA TRANSACCIÓN
        -- =============================================
        COMMIT TRANSACTION;

        SELECT 1 AS success, 'Transferencia registrada correctamente' AS message,
               @mov_id_salida AS mov_id_salida, @mov_codigo_salida AS vale_salida,
               @mov_id_ingreso AS mov_id_ingreso, @mov_codigo_ingreso AS vale_ingreso;

    END TRY
    BEGIN CATCH
        -- Asegurar el Rollback solo si hay una transacción activa
        IF @@TRANCOUNT > 0 
            ROLLBACK TRANSACTION;
            
        -- Limpiar tabla temporal si falló durante su uso
        IF OBJECT_ID('tempdb..#TempDetalle') IS NOT NULL
            DROP TABLE #TempDetalle;

        SELECT 0 AS success, ERROR_MESSAGE() AS message, 
               NULL AS mov_id_salida, NULL AS vale_salida, 
               NULL AS mov_id_ingreso, NULL AS vale_ingreso;
    END CATCH
END
GO""")

    # ==========================================
    # SECCIÓN 5: INGENIERÍA INVERSA CRIPTOGRÁFICA
    # ==========================================
    add_custom_heading("5. Criptografía e Ingeniería Inversa (Contraseñas TECNOTEST)", level=1)
    add_p("Ubicación: backend/config/seguridad.php")
    add_p("La tabla adm_usuario del sistema ERP de escritorio almacena las claves en formato binario bajo una codificación propietaria. Al no existir documentación técnica, se realizó un proceso de ingeniería inversa con reconocimiento de patrones y análisis de muestras controladas.")

    add_custom_heading("5.1. Algoritmo Cifrado César Dinámico por Desplazamiento de Bytes", level=2)
    add_p("Se descubrió que la codificación opera sobre el juego de caracteres Windows-1252 (CP1252) y calcula un desplazamiento dinámico a partir del primer carácter:")
    add_bullet("1. Desplazamiento Base (Shift):", "shift = floor(ord(primer_byte) / 2)")
    add_bullet("2. Transformación Byte a Byte:", "byte_codificado = (byte_original + shift) mod 256")

    add_custom_heading("5.2. Código PHP del Módulo de Seguridad", level=2)
    add_code_block("""<?php
function encodePassword($text) {
    if ($text === '' || $text === null) return '';
    $raw = mb_convert_encoding($text, 'Windows-1252', 'UTF-8');
    $len = strlen($raw);
    if ($len === 0) return '';
    
    $firstByte = ord($raw[0]);
    $shift = intdiv($firstByte, 2);
    
    $encoded = '';
    for ($i = 0; $i < $len; $i++) {
        $b = ord($raw[$i]);
        $encoded .= chr(($b + $shift) % 256);
    }
    return $encoded;
}

function verifyPassword($plainPassword, $storedPassword) {
    $dbClaveRaw = rtrim((string)$storedPassword);
    $inputClaveRaw = encodePassword($plainPassword);
    $inputClaveUtf8 = mb_convert_encoding($inputClaveRaw, 'UTF-8', 'Windows-1252');
    return ($dbClaveRaw === $inputClaveRaw || $dbClaveRaw === $inputClaveUtf8);
}
?>""")

    # ==========================================
    # SECCIÓN 6: BACKEND (MICRO-ENDPOINTS REST)
    # ==========================================
    add_custom_heading("6. Backend: Configuración y Micro-Endpoints REST", level=1)
    add_p("El backend está implementado en PHP 8.2 modular con respuestas estrictas en formato JSON:")

    add_custom_heading("6.1. Servidor Web y .htaccess", level=2)
    add_p("El archivo backend/.htaccess asegura la compatibilidad CORS y el paso de cabeceras:")
    add_code_block("""<IfModule mod_headers.c>
    Header setifempty Access-Control-Allow-Origin "*"
    Header setifempty Access-Control-Allow-Methods "GET, POST, PUT, DELETE, OPTIONS"
    Header setifempty Access-Control-Allow-Headers "Content-Type, Authorization, X-Requested-With"
</IfModule>""")

    add_custom_heading("6.2. Catálogo de Endpoints de la API", level=2)
    add_bullet("POST /api/auth/login.php:", "Autenticación de usuarios, validación de estado activo (usr_status = 1) y generación de token de sesión.")
    add_bullet("GET /api/erp/movimientos.php:", "Listado paginado de notas de movimiento con filtros de almacén, año y mes.")
    add_bullet("GET /api/erp/articulos.php:", "Catálogo y búsqueda paginada de artículos, consulta de lotes con saldo y stock en tiempo real mediante el SP [dbo].[stock_articulo].")
    add_bullet("GET /api/erp/ccostos.php:", "Catálogo de centros de costos activos (mae_ccostomae).")
    add_bullet("GET /api/erp/detalle.php:", "Detalle de artículos de un movimiento específico consultando log_detmov.")
    add_bullet("POST /api/erp/transferencias.php:", "Recepción del payload JSON maestro-detalle y ejecución del procedimiento almacenado sp_registrar_transferencia_almacen.")

    # ==========================================
    # SECCIÓN 7: FRONTEND (ANGULAR 22 + TAILWIND 4.3)
    # ==========================================
    add_custom_heading("7. Frontend: Arquitectura Angular 22 y Seguridad", level=1)
    
    add_custom_heading("7.1. Cifrado AES en Cliente (LocalStorage) y Environments", level=2)
    add_p("Ubicación: frontend/src/app/services/auth.service.ts y frontend/src/environments/environment.ts")
    add_p("Para mitigar el riesgo de manipulación de roles o datos de usuario desde las herramientas de desarrollador del navegador (DevTools), la información de sesión se almacena cifrada mediante el algoritmo simétrico AES de Crypto-JS con una llave secreta corporativa definida en el entorno:")
    add_code_block("""// Encriptación al iniciar sesión
saveSession(user: User, token: string) {
  const encryptedUser = CryptoJS.AES.encrypt(JSON.stringify(user), this.secretKey).toString();
  localStorage.setItem('sigplast_user', encryptedUser);
  localStorage.setItem('sigplast_token', token);
  this.currentUser.set(user);
}""")

    add_custom_heading("7.2. Guards e Interceptor HTTP", level=2)
    add_bullet("userGuard (auth.guard.ts):", "Protege las rutas privadas (/erp, /nuevo-registro). Si no hay sesión válida, redirige al usuario a /login.")
    add_bullet("guestGuard (auth.guard.ts):", "Protege la ruta pública /login. Si el usuario ya está autenticado, lo redirige al panel /erp.")
    add_bullet("authInterceptor (auth.interceptor.ts):", "Inyecta automáticamente la cabecera 'Authorization: Bearer <token>' en cada petición. Atrapa errores 401/403 mostrando alertas automáticas con SweetAlert2 y limpiando la sesión.")

    add_custom_heading("7.3. Estructura de Componentes", level=2)
    add_bullet("LoginComponent:", "Formulario de inicio de sesión con validación reactiva y alternancia de visualización de contraseña.")
    add_bullet("MovimientosComponent:", "Bandeja principal con filtros de almacén/año/mes, paginación persistente y subtablas acordeón para ver el detalle de cada vale bajo demanda.")
    add_bullet("NuevoRegistroComponent:", "Formulario maestro-detalle con escáner de código de barras, modal de búsqueda paginada de artículos con debounce (300ms), selector de lotes con saldo en tiempo real y asignación de centros de costo.")
    add_bullet("NavbarComponent y SidebarComponent:", "Navegación lateral colapsable con cierre automático en dispositivos móviles y visualización de usuario activo.")
    add_bullet("AccessibilityWidgetComponent:", "Panel flotante de control de accesibilidad universal.")
    add_bullet("NotFoundComponent:", "Vista 404 para atrapar rutas inválidas con botón de retorno seguro.")

    # ==========================================
    # SECCIÓN 8: ACCESIBILIDAD Y RESPONSIVIDAD
    # ==========================================
    add_custom_heading("8. Accesibilidad Universal (A11Y) y Diseño Responsivo", level=1)
    
    add_custom_heading("8.1. Funcionalidades de Accesibilidad", level=2)
    add_bullet("Modo Dislexia:", "Aplica la fuente Google Font 'Lexend' con espaciado aumentado de letras (0.05em) y altura de línea 1.6.")
    add_bullet("Filtros de Daltonismo:", "Filtros SVG integrados para Protanopía, Deuteranopía y Tritanopía.")
    add_bullet("Modos de Contraste:", "Modo Monocromático (Escala de grises), Alto Contraste (150%) e Invertido.")
    add_bullet("Escalado de Texto:", "Ajuste dinámico de tamaño de fuente (75%, 100%, 125%, 150%).")
    add_bullet("Modo Noche (Dark Mode):", "Adaptación visual integral a paleta oscura (#0b1329) en todas las vistas, tablas, inputs, scrollbars y popups de SweetAlert2.")

    add_custom_heading("8.2. Optimización para Móviles y Terminales Handhelds", level=2)
    add_p("Ubicación: frontend/src/styles.css")
    add_p("El sistema está optimizado para su uso en terminales colectoras de datos Android / Handhelds industriales:")
    add_bullet("Prevención de Zoom Automático:", "Inputs calibrados a 13px en resoluciones menores a 640px para evitar que el navegador aplique zoom automático al enfocar campos.")
    add_bullet("Desplazamiento Inercial:", "Clases .overflow-x-auto con -webkit-overflow-scrolling: touch para un scroll suave en pantallas táctiles.")
    add_bullet("Zonas de Toque Ampliadas:", "Botones y controles con altura mínima de 44px para facilitar la interacción con guantes de trabajo.")

    doc.add_page_break()

    # ==========================================
    # SECCIÓN 9: MANUAL DE USUARIO PASO A PASO
    # ==========================================
    add_custom_heading("9. Manual de Usuario (Guía Paso a Paso)", level=1)
    add_p("Esta sección describe el uso operativo del sistema con espacios listos para incorporar las capturas de pantalla de la aplicación:")

    add_custom_heading("9.1. Inicio de Sesión (Login)", level=2)
    add_p("1. Ingrese a la URL de la aplicación en el navegador.")
    add_p("2. Escriba su código de Usuario (ej. ADMINISTRA) y Contraseña corporativa.")
    add_p("3. Presione el botón 'Iniciar Sesión'. El sistema validará su estado en SQL Server y lo redirigirá a la bandeja principal.")
    add_screenshot_placeholder("Pantalla de Inicio de Sesión (Login)", "Captura de la vista de login con los campos de usuario, contraseña y logo institucional")

    add_custom_heading("9.2. Consulta y Filtrado de Movimientos (Bandeja ERP)", level=2)
    add_p("1. Al ingresar a '/erp', visualizará la bandeja de notas de almacén.")
    add_p("2. Seleccione el Almacén deseado (ej. '001 - ALMACEN VERDE MP'), el Año y el Mes.")
    add_p("3. Utilice los controles de paginación para navegar entre las páginas o ajustar la cantidad de filas por vista (10, 20, 50, 100).")
    add_screenshot_placeholder("Bandeja Principal de Movimientos", "Captura de la grilla de notas de salida con selectores de almacén, año, mes y paginación")

    add_custom_heading("9.3. Visualización de Detalles de un Vale", level=2)
    add_p("1. Haga clic sobre cualquier fila de movimiento en la tabla.")
    add_p("2. Se desplegará un acordeón animado consultando en tiempo real los artículos, lotes, cantidades transferidas y unidad de medida.")
    add_screenshot_placeholder("Detalle de Movimiento Expandido", "Captura del acordeón expandido mostrando la subtabla con los artículos y lotes del vale")

    add_custom_heading("9.4. Registro de una Nueva Transferencia de Almacén", level=2)
    add_p("1. Presione el botón '+ Nuevo Registro' ubicado en la esquina superior derecha de la bandeja o en el menú lateral.")
    add_p("2. Seleccione el Almacén Origen y el Tipo de Movimiento (102 - Salida por Transferencia).")
    add_p("3. Al detectar que es una transferencia, el sistema habilitará automáticamente el selector de Almacén Destino.")
    add_p("4. Ingrese la Fecha de Emisión y la Glosa (observación) del movimiento.")
    add_screenshot_placeholder("Formulario de Cabecera de Transferencia", "Captura del formulario maestro con almacén origen, destino, tipo de movimiento y glosa")

    add_custom_heading("9.5. Búsqueda y Selección de Artículos", level=2)
    add_p("Existen dos formas de agregar artículos a la transferencia:")
    add_bullet("Lectura por Código de Barras:", "Enfoque el campo 'Escanear Código de Barras' y dispare la pistola lectora EAN/Code128.")
    add_bullet("Búsqueda Manual Paginada:", "Haga clic en 'Buscar Artículo' para abrir el modal interactivo. Escriba el nombre o código para filtrar con búsqueda reactiva y presione 'Seleccionar'.")
    add_screenshot_placeholder("Modal de Búsqueda de Artículos", "Captura del modal de catálogo de artículos con barra de búsqueda y paginación")

    add_custom_heading("9.6. Selección de Lotes y Centro de Costos", level=2)
    add_p("1. En la grilla de ítems agregados, presione sobre el botón del Lote para abrir el modal de lotes.")
    add_p("2. El sistema consultará con el Stored Procedure únicamente los lotes con saldo disponible.")
    add_p("3. Seleccione el lote deseado, digite la cantidad a transferir y asigne el Centro de Costos correspondiente.")
    add_screenshot_placeholder("Modal de Selección de Lotes con Saldo", "Captura del modal de lotes mostrando los saldos físicos disponibles por lote")

    add_custom_heading("9.7. Emisión y Confirmación de la Transferencia", level=2)
    add_p("1. Revise los ítems en la tabla resumen.")
    add_p("2. Presione el botón verde 'Guardar Transferencia'.")
    add_p("3. El sistema ejecutará el procedimiento almacenado en SQL Server y SweetAlert2 le mostrará una alerta de éxito confirmando el número de Vale de Salida generado y el Vale de Ingreso en el almacén de destino.")
    add_screenshot_placeholder("Alerta SweetAlert2 de Transferencia Exitosa", "Captura del popup modal de SweetAlert2 confirmando la creación de los vales")

    add_custom_heading("9.8. Panel de Accesibilidad y Modo Noche", level=2)
    add_p("1. Presione el botón flotante o icono de accesibilidad en la barra de navegación.")
    add_p("2. Ajuste el tamaño de letra, active filtros de daltonismo o active el Modo Noche.")
    add_screenshot_placeholder("Panel de Accesibilidad y Modo Oscuro", "Captura del panel desplegable de accesibilidad y la interfaz en modo oscuro")

    doc.add_page_break()

    # ==========================================
    # SECCIÓN 10: GUÍA DE DESPLIEGUE LOCAL
    # ==========================================
    add_custom_heading("10. Guía de Despliegue y Configuración Local", level=1)
    add_p("Esta sección detalla los pasos y requisitos indispensables para ejecutar el proyecto completo en un entorno de desarrollo local:")

    add_custom_heading("10.1. Requisitos Previos de Software", level=2)
    add_bullet("Sistema Operativo:", "Windows 10 / Windows 11 de 64 bits.")
    add_bullet("Servidor Web y PHP:", "XAMPP con PHP 8.2 instalado en C:\\xampp.")
    add_bullet("Drivers Microsoft SQL Server para PHP:", "Extensiones php_sqlsrv y php_pdo_sqlsrv (versión 8.2 Thread Safe x64).")
    add_bullet("Microsoft ODBC Driver:", "ODBC Driver 18 for SQL Server (o Driver 17) instalado en Windows.")
    add_bullet("Node.js y npm:", "Node.js versión 20+ o 24+ y npm versión 10+.")
    add_bullet("Angular CLI:", "Angular CLI versión 22 instalada globalmente (npm install -g @angular/cli).")
    add_bullet("Servidor SQL Server:", "Acceso a la instancia de SQL Server 2022 (ej. 192.168.1.3\\SQLEXPRESS) con la base de datos TECNOTEST.")

    add_custom_heading("10.2. Configuración de Extensiones en PHP (XAMPP)", level=2)
    add_p("1. Descargue los drivers de Microsoft SQL Server para PHP 8.2.")
    add_p("2. Copie los archivos 'php_sqlsrv_82_ts_x64.dll' y 'php_pdo_sqlsrv_82_ts_x64.dll' en la carpeta 'C:\\xampp\\php\\ext\\'.")
    add_p("3. Abra el archivo 'C:\\xampp\\php\\php.ini' y agregue las siguientes líneas:")
    add_code_block("""extension=php_sqlsrv_82_ts_x64.dll
extension=php_pdo_sqlsrv_82_ts_x64.dll""")
    add_p("4. Reinicie el servicio Apache desde el panel de control de XAMPP.")

    add_custom_heading("10.3. Ubicación del Backend en XAMPP", level=2)
    add_p("Coloque la carpeta del proyecto dentro del directorio web de Apache:")
    add_code_block("""C:\\xampp\\htdocs\\transferencia-almacenes\\backend\\""")

    add_custom_heading("10.4. Configuración de Conexión a Base de Datos", level=2)
    add_p("Edite el archivo backend/config/conexion.php para verificar las credenciales de SQL Server:")
    add_code_block("""$serverName = "192.168.1.3\\\\SQLEXPRESS";
$connectionInfo = array(
    "Database" => "TECNOTEST",
    "UID" => "sig",
    "PWD" => "Sig2025$$",
    "CharacterSet" => "UTF-8",
    "TrustServerCertificate" => true
);
$conn = sqlsrv_connect($serverName, $connectionInfo);""")

    add_custom_heading("10.5. Instalación y Ejecución del Frontend (Angular)", level=2)
    add_p("1. Abra una terminal de comandos (PowerShell o CMD) y navegue a la carpeta frontend:")
    add_code_block("""cd "C:\\Users\\manue\\OneDrive\\Desktop\\TRABAJO\\PAGINAS\\PRODUCTO FINAL\\transferencia-almacenes\\frontend"
npm install""")

    add_p("2. Inicie el servidor de desarrollo de Angular:")
    add_code_block("""npm start
# O alternativamente:
ng serve --host 0.0.0.0 --port 4200""")

    add_p("3. Abra su navegador web e ingrese a:")
    add_code_block("""http://localhost:4200""")

    add_callout_box("Nota de Acceso en Red Local (Móviles y Handhelds)", 
                    "Para acceder desde terminales móviles o colectoras conectadas a la misma red WiFi, reemplace 'localhost' por la dirección IP local de su computadora (ej. http://192.168.1.50:4200). El backend resolverá automáticamente la URL base gracias a BaseApiService.")

    # Guardar documento
    output_path = r"c:\Users\manue\OneDrive\Desktop\TRABAJO\PAGINAS\PRODUCTO FINAL\transferencia-almacenes\DOCUMENTACION_SISTEMA_TRANSFERENCIA_ALMACENES_SIGPLAST.docx"
    doc.save(output_path)
    print(f"Documento generado exitosamente en: {output_path}")

if __name__ == "__main__":
    create_document()
